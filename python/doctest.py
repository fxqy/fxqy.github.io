#!/usr/bin/python3
 
from docx import Document
from docx.shared import Inches,Pt
import psutil

doc=Document()

doc.add_heading("windows本机进程",0)
doc.add_heading("获取进程并写入Word文档",level=1)
doc.add_paragraph("插入图片")#.add_paragraph('Intense quote', style='IntenseQuote') style='ListBullet' style='ListNumber'
doc.add_picture('111.png',width=Inches(6.0),height=Inches(3.0))

processes = [(p.pid, p.name()) for p in psutil.process_iter()]

table=doc.add_table(rows=1,cols=3)
table.style=doc.styles['Table Grid']

head_cells = table.rows[0].cells
head_cells[0].text = '#'
head_cells[1].text = '进程ID'
head_cells[2].text = '名称'

for i,proc in enumerate(processes):
    cells = table.add_row().cells
    cells[0].text = str(i+1)
    cells[1].text = str(proc[0])
    cells[2].text = proc[1]
    
#doc.add_page_break()
doc.save('C:\\Users\\Dong\\Desktop\\pydoc.docx')
